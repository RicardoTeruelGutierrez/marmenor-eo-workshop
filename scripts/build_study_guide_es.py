"""Genera la guía del instructor / guión didáctico en ESPAÑOL (.docx).

Guía combinada, sección a sección, de los tres notebooks: para cada parte,
"en una línea", "qué decir", "por qué importa" y "preguntas probables".
Ejecutar:  python scripts/build_study_guide_es.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).resolve().parent.parent / "Guia_de_Estudio_Instructor.docx"

INK   = RGBColor(0x0E, 0x1A, 0x26)
ACCENT= RGBColor(0xC2, 0x41, 0x0C)
TEAL  = RGBColor(0x11, 0x5E, 0x59)
GREY  = RGBColor(0x6B, 0x61, 0x57)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for sec in doc.sections:
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)


def _style_heading(level, size, color, before, after):
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Georgia"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

_style_heading(1, 18, ACCENT, 18, 8)
_style_heading(2, 13.5, TEAL, 14, 4)


def h1(text, page_break=True):
    p = doc.add_heading(text, level=1)
    if page_break: p.paragraph_format.page_break_before = True
    return p

def h2(text): return doc.add_heading(text, level=2)

def body(text, italic=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def labeled(label, text, label_color=ACCENT):
    p = doc.add_paragraph()
    lr = p.add_run(label + " "); lr.bold = True; lr.font.color.rgb = label_color
    p.add_run(text)
    return p

def qa(pairs):
    p = doc.add_paragraph(); r = p.add_run("Preguntas probables"); r.bold = True; r.font.color.rgb = TEAL
    for q, a in pairs:
        pq = doc.add_paragraph(); pq.paragraph_format.space_after = Pt(2)
        rq = pq.add_run("P.  "); rq.bold = True; rq.font.color.rgb = ACCENT
        rq2 = pq.add_run(q); rq2.bold = True
        pa = doc.add_paragraph(); pa.paragraph_format.left_indent = Inches(0.3)
        pa.paragraph_format.space_after = Pt(8)
        ra = pa.add_run("R.  "); ra.bold = True
        pa.add_run(a)

def section(title, oneline, say, why, questions=None):
    h2(title)
    labeled("En una línea:", oneline)
    labeled("Qué decir:", say, label_color=TEAL)
    labeled("Por qué importa:", why, label_color=TEAL)
    if questions: qa(questions)


# ── TÍTULO ──────────────────────────────────────────────────────────────
t = doc.add_paragraph()
r = t.add_run("Mar Menor — Taller de Observación de la Tierra")
r.font.name = "Georgia"; r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = INK
sub = doc.add_paragraph()
rs = sub.add_run("Guía del Instructor y Guión Didáctico")
rs.font.name = "Georgia"; rs.font.size = Pt(16); rs.italic = True; rs.font.color.rgb = ACCENT
meta = doc.add_paragraph()
rm = meta.add_run("Qué decir, por qué importa y las preguntas que te harán — "
                  "sección a sección en los tres notebooks.")
rm.font.size = Pt(12); rm.font.color.rgb = GREY

doc.add_paragraph()
labeled("Cómo usar esta guía:",
        "Cada sección de los notebooks tiene cuatro partes — un resumen en una "
        "línea, una explicación lista para decir en voz alta, la razón por la que "
        "importa, y las preguntas que suelen hacer los alumnos con respuestas "
        "modelo. Léela una vez de principio a fin para dominar el relato; en clase "
        "basta con ojear las etiquetas en negrita. Combínala con las "
        "presentaciones (slides/module1–3) y los notebooks en vivo.")
labeled("El hilo conductor:", "Tres notebooks, una laguna. El Módulo 1 lee la "
        "laguna desde el espacio; el Módulo 2 lo valida contra las boyas y añade "
        "machine learning; el Módulo 3 enseña a los alumnos a construir ellos "
        "mismos toda la capa de datos. Todo gira en torno a las crisis "
        "documentadas del Mar Menor (2016, DANA de 2019, mortandad de 2021, "
        "rebrote de 2025).")

# ── PARTE 0 ─────────────────────────────────────────────────────────────
h1("Parte 0 — La historia que estás contando")
body("El Mar Menor es la mayor laguna costera hipersalina de Europa: ~135 km², "
     "profundidad máxima 7 m, separada del Mediterráneo por La Manga, un cordón "
     "litoral de 22 km. Décadas de nitratos de la agricultura del Campo de "
     "Cartagena, más episodios de DANA y el calentamiento del Mediterráneo, la "
     "han llevado a crisis de eutrofización repetidas. Cada colapso es visible "
     "desde órbita a 10 metros de resolución — por eso es un caso de estudio "
     "ideal.")
labeled("El arco que hay que tener en la cabeza:",
        "presión (nutrientes) → disparador (calor / inundación) → bloom "
        "(clorofila) → colapso de oxígeno (anoxia) → mortandad de peces. Cada "
        "módulo ilumina una parte de esta cadena con una fuente de datos distinta.")
labeled("Un marco honesto para repetir:",
        "la observación de la Tierra puede acortar el ciclo de la señal a la "
        "decisión, pero la variable lenta es la política. Cada algoritmo de aquí "
        "ya corre en algún pipeline en Murcia; la laguna sigue degradándose. "
        "Mantén esa tensión — mantiene el trabajo técnico con los pies en la tierra.")

# ── MÓDULO 1 ────────────────────────────────────────────────────────────
h1("Módulo 1 — Extracción de datos satelitales")
body("Objetivo: pasar de la reflectancia bruta de Sentinel a las huellas "
     "visibles del colapso. El notebook se divide en tres secciones — A "
     "(sintética, corre sin conexión), B (el flujo real CDSE/STAC, mostrado pero "
     "no ejecutado), C (figuras de escenas Sentinel-2 reales descargadas).")

section("Setup, chequeo de dependencias y la división A/B/C",
    "El notebook falla pronto si falta un paquete o un archivo, y luego declara "
    "su estructura de tres secciones.",
    "Empieza señalando la tabla de la cabecera: la Sección A corre sin internet, "
    "la B necesita una cuenta gratuita, la C necesita los GeoTIFF reales (y se "
    "auto-omite si no están). Ejecuta la celda de setup; comprueba "
    "numpy/pandas/rasterio, etc., y carga el contorno de la laguna.",
    "Responde directamente a una crítica de que los datos sintéticos y reales se "
    "mezclaban. Los chequeos tempranos evitan que un alumno depure un error "
    "críptico a mitad de camino.",
    [("¿Por qué no un único notebook grande?",
      "Porque «lo que corre en cualquier sitio» y «lo que necesita la nube» "
      "tienen modos de fallo distintos. Separarlos mantiene la clase avanzando "
      "aunque no haya Wi-Fi ni cuenta.")])

section("A.1 — Aguas Case-2: por qué fallan los algoritmos estándar",
    "La óptica del Mar Menor está dominada por sedimentos, CDOM y macrófitos, no "
    "solo por el fitoplancton — así que los algoritmos oceánicos de clorofila "
    "fallan.",
    "Explica Case-1 frente a Case-2. Case-1 (océano abierto) supone que solo el "
    "fitoplancton controla el color, así que un cociente azul/verde funciona. El "
    "Mar Menor es Case-2: el sedimento en suspensión sube la reflectancia "
    "azul-verde, el CDOM (materia orgánica disuelta coloreada de la escorrentía) "
    "absorbe en el azul, y los macrófitos del fondo añaden señal en el "
    "infrarrojo cercano. Cada uno mueve el espectro independientemente de la "
    "clorofila. La solución es la corrección atmosférica C2RCC (red neuronal) y "
    "luego un índice calibrado localmente. Nuestro proxy es el cociente "
    "red-edge Rrs(B05)/Rrs(B04).",
    "Es la columna vertebral conceptual de toda la parte satelital. Si los "
    "alumnos no la captan, se fiarán demasiado de cualquier cociente de bandas.",
    [("¿Sentinel-2 L2A es lo mismo que C2RCC?",
      "No — y hay que decirlo claro. L2A (Sen2Cor) es una corrección atmosférica "
      "de TIERRA. Para trabajo riguroso en agua se ejecuta C2RCC (p. ej. en SNAP "
      "o con un evalscript de Sentinel Hub). Nuestros valores cacheados solo "
      "imitan un producto tipo C2RCC."),
     ("¿Por qué el red-edge (B05) y no el clásico azul/verde?",
      "La clorofila absorbe fuertemente en el rojo (B04, 665 nm) y refleja en el "
      "red-edge (B05, 705 nm). El cociente sube con la concentración y es mucho "
      "más robusto al sedimento que el azul/verde en agua turbia.")])

section("A.3 — La serie temporal de clorofila de toda la laguna",
    "Promediar los puntos de muestreo a una media semanal reproduce las cuatro "
    "crisis documentadas en los meses correctos.",
    "Muestra la curva con las ventanas de 2016, 2019, 2021 y 2025 sombreadas. "
    "Recalca que los datos sintéticos están calibrados con la literatura "
    "publicada, así que las crisis aparecen en las fechas y magnitudes correctas "
    "— es un sustituto docente de lo que produciría un pipeline real de varios "
    "años.",
    "Ancla el resto del módulo en eventos reales y fechados, y establece 2021 "
    "como la crisis protagonista.",
    [("Son sintéticos — ¿no socava eso la lección?",
      "No, mientras seamos explícitos. La estadística está calibrada con la "
      "literatura; los métodos (resampleo, lógica de anomalías) son exactamente "
      "los que ejecutarías con datos reales.")])

section("A.6 — La ola de calor de 2021 y su vínculo con el oxígeno",
    "Una ola de calor de +2 °C baja la solubilidad del oxígeno ~0.3 mg/L; con el "
    "bloom respirando de noche y la estratificación sellando el fondo, eso "
    "provoca la anoxia.",
    "Recorre la cadena causal de forma cuantitativa. Agosto de 2021 estuvo a "
    "~30.5 °C frente a ~27.5 °C de un año normal — una anomalía de +1.9 °C. Con "
    "la curva de solubilidad de Benson–Krause, ese calentamiento por sí solo "
    "elimina ~0.3 mg/L de la capacidad de oxígeno del agua. De noche el bloom "
    "consume oxígeno; la estratificación impide que el fondo se reoxigene. "
    "Súmalo y obtienes la hipoxia tras la mortandad. Mostramos el campo de SST "
    "recortado al contorno real de la laguna, su anomalía y la curva de "
    "solubilidad, lado a lado.",
    "Convierte la SST de «una capa extra simpática» en un motor computable del "
    "colapso ecológico — el pico emocional y científico del Módulo 1.",
    [("¿Por qué recortar la SST al contorno de la laguna?",
      "Versiones anteriores mostraban píxeles cayendo sobre tierra, lo que se "
      "veía mal y diluía la señal. Recortar al contorno derivado del NIR mantiene "
      "el estadístico honesto y la figura limpia."),
     ("¿0.3 mg/L es de verdad suficiente?",
      "Por sí solo, no — pero se suma a la demanda biológica de oxígeno nocturna "
      "y a la estratificación. La lección es que varios estresores moderados se "
      "combinan de forma no lineal en un evento letal.")])

section("Sección B — El flujo real CDSE / STAC / COG",
    "Cómo descubrir y descargar datos Sentinel reales hoy, mostrado como patrones "
    "(no ejecutado en clase).",
    "Insiste en tres correcciones frente a tutoriales antiguos. (1) SciHub está "
    "retirado; el punto de entrada es el Copernicus Data Space Ecosystem (CDSE). "
    "(2) El endpoint STAC actual es stac.dataspace.copernicus.eu/v1 y la "
    "colección es en minúsculas sentinel-2-l2a (el endpoint antiguo quedó "
    "deprecado en nov. 2025). (3) Usa L2A (reflectancia de superficie), no L1C "
    "(tope de atmósfera). Luego el truco de escalabilidad: un Cloud-Optimised "
    "GeoTIFF permite leer solo la ventana del AOI — unos pocos MB en vez de una "
    "escena de ~800 MB.",
    "Son los detalles exactos que hacen que el código del alumno funcione o falle "
    "en 2026. La lectura de ventana del COG es el mejor hábito para escalar.",
    [("¿Por qué no ejecutamos las llamadas autenticadas en clase?",
      "Las credenciales son personales y la cuota es limitada. Mostramos el "
      "patrón y ejecutamos el mirror de AWS sin login, para que todos lo logren.")])

section("Sección C — Figuras Sentinel-2 reales (True Colour, NDCI, zonas)",
    "Cuatro escenas reales sin nubes dan True-Colour pixel-perfecto, un mapa NDCI "
    "de clorofila por píxel y un análisis por zonas Norte/Centro/Sur.",
    "Recalca dos aciertos metodológicos. (1) Selección de escenas: medimos la "
    "cobertura válida Y las nubes sobre la propia laguna, no el % de nubes de "
    "toda la tesela del catálogo. Por eso se descartó la escena del 14-jul-2021 "
    "— solo rozaba un borde de pasada (5.5 % de cobertura). (2) El contorno de la "
    "laguna se extrae de la banda infrarroja cercana (el agua absorbe el NIR), "
    "dando 191 vértices exactos en lugar de un polígono dibujado a mano. El NDCI "
    "va de ≈ −0.28 en agua limpia de invierno a ≈ −0.05 en el bloom de verano; "
    "las barras por zona muestran la costa noroeste (agrícola) como la más alta.",
    "Aquí el índice abstracto se convierte en un mapa creíble de un lugar real, y "
    "donde los alumnos aprenden que las decisiones de limpieza cambian la ciencia.",
    [("¿Por qué el NDCI sale casi todo negativo?",
      "El NDCI absoluto depende de la corrección atmosférica y de las bandas "
      "exactas; lo que importa es el salto relativo (~+0.25 de invierno a bloom) "
      "y el patrón espacial, ambos robustos."),
     ("¿Podríamos calcular clorofila real en mg/m³?",
      "Sí, con un C2RCC calibrado localmente + regresión. Dejamos las escenas "
      "reales como NDCI y usamos la serie sintética para la tendencia en mg/m³, "
      "para ser honestos sobre qué admite cada fuente.")])

# ── MÓDULO 2 ────────────────────────────────────────────────────────────
h1("Módulo 2 — Integración in-situ y Machine Learning")
body("Objetivo: conectar el registro satelital con la verdad de campo, "
     "validarlo, y aprender a predecir y alertar de las crisis — con la "
     "honestidad estadística que separa una demo docente de un resultado inflado.")

section("Sección 1 — Las redes de monitoreo in-situ",
    "Datos diarios de boyas (CARM/IMIDA) más hidrología de cuenca (CHS-SAIH) son "
    "la verdad de campo y la presión aguas arriba.",
    "Presenta las cuatro fuentes: sondas en tiempo real de CARM (T, S, chl-a, "
    "OD, turbidez), agro-meteorología de IMIDA (la presión de nutrientes), "
    "hidrología de cuenca CHS-SAIH (la señal de inundación) y Copernicus Marine "
    "INS-TAC. Nuestro archivo cacheado imita las exportaciones de CARM: cinco "
    "estaciones, diario, 2016–2025.",
    "Los satélites necesitan verdad de campo para ser creíbles; y los datos "
    "agrícolas aguas arriba son lo que en última instancia degrada la laguna.",
    None)

section("Sección 1.1–1.2 — Las crisis en la señal bruta",
    "Antes de modelar, dibuja la hipoxia de 2021 y la DANA de 2019 directamente "
    "desde las boyas.",
    "Para 2021, dibuja el oxígeno disuelto contra la línea de hipoxia de 2 mg/L "
    "con la SST superpuesta — es el gráfico que acompañó los titulares de la "
    "mortandad. Para 2019, muestra la DANA como una cascada: 300+ mm de lluvia "
    "en 48 h → la salinidad baja ~4 PSU → el nitrato sube diez veces.",
    "Crea el hábito de mirar los eventos antes de fiarse de cualquier modelo "
    "sobre ellos, y hace concreta la detección de anomalías posterior.",
    None)

section("Sección 2 — Match-ups: validar el satélite",
    "Co-localiza cada punto Sentinel-2 con su boya más cercana (4 km, mismo día) "
    "y puntúa la estimación con métricas adecuadas.",
    "Explica la convención de match-up (el color del océano usa 3 km / 3 h; "
    "nosotros 4 km / mismo día porque los datos in-situ son diarios). Reporta R², "
    "RMSE y RMSLE (forma logarítmica, porque la clorofila es aprox. log-normal) y "
    "el sesgo, y dibuja siempre la recta 1:1. El diagrama de dispersión contra la "
    "1:1 es la prueba visual honesta.",
    "Ninguna afirmación de habilidad es creíble sin un match-up. Es la disciplina "
    "que separa un mapa bonito de un producto validado.",
    [("¿Por qué transformar el error en logaritmo (RMSLE)?",
      "La clorofila abarca órdenes de magnitud y es aprox. log-normal; un RMSE "
      "absoluto lo dominan unos pocos puntos de bloom alto, ocultando el "
      "rendimiento en concentraciones típicas.")])

section("Sección 3 — Estimación ML y validación cruzada HONESTA",
    "Un regresor de gradient boosting predice chl-a desde reflectancias; lo "
    "validamos de dos formas que responden a preguntas distintas.",
    "Primero, TimeSeriesSplit respeta el orden cronológico y responde «¿podemos "
    "predecir fechas futuras?» — y sale fuerte. Luego GroupKFold con la estación "
    "como grupo deja fuera ESTACIONES ENTERAS y responde «¿podemos predecir un "
    "sitio no visto y no monitorizado?» — y el R² cae a ~0.17. La diferencia es "
    "la lección: un modelo puede verse genial en el tiempo y ser débil en el "
    "espacio. Reporta el número espacial, más difícil, cuando hables de desplegar "
    "en un sitio nuevo.",
    "Es el titular del módulo en rigor — corrige directamente la crítica de «el "
    "ML parece demasiado optimista» y enseña un hábito transferible.",
    [("¿Por qué el R² espacial es tan bajo? ¿El modelo es malo?",
      "No malo — evaluado con honestidad. Predecir una estación que el modelo "
      "nunca vio es genuinamente más difícil que predecir una fecha nueva en una "
      "estación conocida. El número bajo es la estimación de despliegue real."),
     ("¿Qué número pongo en un paper?",
      "Ambos, bien etiquetados. Pero el que refleja el uso real en un sitio nuevo "
      "es el de GroupKFold (espacial).")])

section("Sección 4 — Detección multivariante de anomalías (alerta temprana)",
    "Un Isolation Forest entrenado SOLO con los años tranquilos 2017–2018, y "
    "aplicado fuera de muestra, marca anomalías conjuntas en seis variables.",
    "Haz central el punto de la ventana de entrenamiento. Si ajustas el detector "
    "con toda la serie, aprende las crisis como «normales» y la alerta temprana "
    "se vuelve circular. Lo ajustamos con la base tranquila 2017–2018 y lo "
    "aplicamos hacia delante a 2019+, así toda marca posterior es genuinamente "
    "fuera de muestra. Vigila chl-a, turbidez, OD, SST, salinidad y nitrato de "
    "forma conjunta — las anomalías conjuntas captan precursores que un umbral "
    "de chl-a único se pierde.",
    "Es la diferencia entre un detector que «predice» el pasado y uno que de "
    "verdad podría alertar a una administración.",
    [("¿Por qué Isolation Forest y no un umbral simple?",
      "Un umbral univariante se pierde las combinaciones — p. ej. chl-a "
      "moderadamente alta + OD bajo + SST alta a la vez. Isolation Forest puntúa "
      "el estado multivariante conjunto.")])

section("Sección 4.1 — Lead time frente a retraso de detección",
    "Separa una marca ANTES de una crisis (lead time, un acierto) de una marca "
    "DESPUÉS (retraso de detección, un fallo) — nunca las mezcles.",
    "Muestra la tabla: 2019 se marcó 42 días antes (alerta genuina); 2021 y 2025 "
    "solo se captaron 1–2 días tras el inicio (retraso de detección, es decir, "
    "fallos). El código ingenuo barría una ventana que empezaba antes y acababa "
    "después, tomaba la primera marca y la llamaba «lead time» aunque llegara "
    "tarde. Separamos las dos métricas explícitamente.",
    "Es un arreglo sutil de contabilidad que evita que los papers de monitoreo se "
    "vendan de más — un gran momento docente sobre integridad científica.",
    [("Entonces, ¿nuestro detector falla casi siempre?",
      "Acierta espectacularmente una vez (2019) y llega tarde dos. Ese "
      "boletín honesto es el objetivo — y motiva añadir variables de la cuenca "
      "para adelantar el lead time, que es el ejercicio para casa.")])

# ── MÓDULO 3 ────────────────────────────────────────────────────────────
h1("Módulo 3 — Construye tu propia base de datos EO")
body("Objetivo: que los alumnos construyan ellos mismos toda la capa de datos y "
     "ejecuten descargas reales en directo. Enmarca el módulo con la metáfora de "
     "la cocina: la compra → lavar → cortar → la nevera → cocinar → llevar el "
     "restaurante.")

section("Resumen Parte 1 — el pipeline y el mapa de fuentes",
    "Seis pasos (descubrir, limpiar, extraer, almacenar, analizar, "
    "industrializar) y un mapa de ocho fuentes gratuitas, cuatro sin login.",
    "Da el mapa completo antes de los detalles. Luego la regla práctica: "
    "prototipa hoy con las cuatro fuentes sin login (NASA GIBS, COGs Sentinel-2 "
    "de AWS, Open-Meteo Air-Quality y Marine); pasa a las de cuenta (CDSE, "
    "Copernicus Marine, NASA Earthdata) cuando necesites un producto concreto o "
    "el archivo histórico completo.",
    "Los alumnos se van sabiendo no un solo truco de descarga, sino dónde "
    "conseguir casi cualquier dato ambiental gratis — el objetivo real del módulo.",
    None)

section("Secciones 1.1–1.4 — Descargas en directo (ejecútalas en clase)",
    "Cuatro descargas reales sin credenciales que los alumnos ven llegar: una "
    "imagen GIBS, una ventana de COG Sentinel-2 y dos series de Open-Meteo.",
    "Ejecuta cada celda en vivo. GIBS: una URL devuelve un PNG true-colour para "
    "cualquier fecha desde el año 2000 — cambia la fecha y reejecuta. COG "
    "Sentinel-2: busca en el catálogo, abre la cabecera del COG SIN descargar, y "
    "luego lee solo la ventana del AOI — unos 8 MB, ~3 % de la escena. Open-Meteo: "
    "un GET devuelve JSON que se vuelve un DataFrame con PM2.5/PM10/ozono "
    "(¡contaminación!) y temperatura superficial del mar por horas. Cada celda "
    "degrada con elegancia si no hay conexión.",
    "Es el momento en que «la EO deja de ser abstracta» — los alumnos ven llegar "
    "bytes reales y entienden que pueden hacerlo para su propia costa hoy.",
    [("¿De verdad el COG descarga solo parte del archivo?",
      "Sí. Un Cloud-Optimised GeoTIFF está internamente en mosaicos e indexado; "
      "rasterio pide por HTTP solo los mosaicos que solapan tu ventana. Por eso "
      "es ~3 % de los bytes."),
     ("¿Por qué AWS en vez de Copernicus para la demo en vivo?",
      "El mirror Open-Data de AWS no necesita login, así que todos lo logran al "
      "instante. Copernicus viene después, para lo que desbloquea la cuenta.")])

section("Sección 1.5 — Una cuenta gratuita de Copernicus y qué desbloquea",
    "Registrarse lleva ~2 minutos; un cliente OAuth habilita el procesado en "
    "servidor que el mirror de AWS no puede dar.",
    "Recorre los tres pasos de registro (regístrate en dataspace.copernicus.eu → "
    "crea un cliente OAuth en el dashboard de Sentinel Hub → guarda id/secret "
    "como variables de entorno). Luego los dos desbloqueos, ambos verificados en "
    "vivo: la Process API manda un evalscript diminuto y devuelve un producto "
    "terminado (mapa NDCI, imagen true-colour o pila de bandas crudas) para "
    "cualquier fecha; la Statistical API devuelve una serie temporal completa del "
    "AOI en una sola petición sin descargar NINGUNA imagen. Muestra la firma "
    "espectral de las bandas crudas — la reflectancia subiendo hacia el red-edge "
    "es la huella de la clorofila.",
    "Muestra el salto de «descarga píxeles y calcula tú» a «pide al servidor "
    "justo el producto que quieres» — el flujo profesional.",
    [("¿Cuánto cuesta la cuenta?",
      "Nada — gratis, sin tarjeta, ~30.000 unidades de procesado al mes, de "
      "sobra para un proyecto de doctorado."),
     ("¿Qué es un evalscript?",
      "Una pequeña función JavaScript que le dice a Sentinel Hub qué debe ser "
      "cada píxel de salida. Cambia la función y cambias el producto — NDCI, true "
      "colour o bandas crudas — sin tocar el código de descarga.")])

section("Secciones 2–4 — Limpiar, extraer y almacenar (SQLite + Parquet)",
    "Convierte píxeles en filas ordenadas y fiables, y guárdalas en una base de "
    "datos tipada, indexada e idempotente más Parquet para los arrays masivos.",
    "Limpiar: enmascara lo que no es agua y las nubes, mide la cobertura. "
    "Extraer: calcula el NDCI por escena y por zona en filas de formato largo "
    "(fecha, fuente, variable, zona, valor). Almacenar: un esquema SQLite con "
    "claves primarias, índices y — crucial — una restricción UNIQUE para que "
    "INSERT OR REPLACE haga las reejecuciones idempotentes (sin duplicados). "
    "Mantén los millones de valores por píxel en Parquet, referenciados desde la "
    "base de datos, no metidos en filas SQL.",
    "Es el corazón de ingeniería: por qué una base de datos supera a un montón de "
    "CSVs, y por qué la idempotencia es lo que permite que un pipeline corra en "
    "un cron para siempre.",
    [("¿Por qué SQLite y no PostgreSQL?",
      "SQLite es un solo archivo, sin instalación, con SQL completo — perfecto "
      "para un proyecto de doctorado. Mostramos la ruta de mejora a PostGIS/"
      "DuckDB para cuando se quede pequeño."),
     ("¿Qué me aporta «idempotente»?",
      "Puedes reprocesar cualquier fecha y la base actualiza esa fila en vez de "
      "añadir un duplicado — esencial para un trabajo nocturno que pueda solaparse "
      "con ejecuciones previas.")])

section("Sección 5 — El estudio de contaminación, directo desde SQL",
    "Consultas SQL cortas producen una tendencia mensual, un recuento de "
    "superaciones de «mal estado», comparaciones por zona y un join "
    "satélite↔in-situ.",
    "Demuestra que las preguntas que hace una administración son ahora una "
    "consulta cada una: clorofila media mensual, una tendencia robusta de "
    "Theil-Sen (resistente a los picos de bloom), el porcentaje de observaciones "
    "por encima de un umbral de «mal estado» por año, y un JOIN de satélite e "
    "in-situ. Los recuentos de superación se disparan exactamente en 2019 y 2021 "
    "— la base de datos detecta las crisis documentadas por sí sola.",
    "Cierra el círculo: la capa de datos que construyeron responde preguntas "
    "reales de contaminación, probando que todo el pipeline funciona de extremo a "
    "extremo.",
    [("¿Por qué Theil-Sen en vez de mínimos cuadrados?",
      "Theil-Sen es robusto a valores atípicos — unos pocos meses de bloom "
      "extremo no arrastran la recta como en mínimos cuadrados ordinarios.")])

section("Sección 6 — Industrializar y escalar",
    "Envuelve los pasos 1–4 en una función idempotente, prográmala, y conoce la "
    "ruta de mejora a PostGIS / DuckDB.",
    "Muestra el esqueleto de update_database(): descubrir → descargar ventana → "
    "limpiar → extraer → upsert, todo idempotente. Luego la tabla de escalado: "
    "PostGIS para consultas espaciales y muchos usuarios, DuckDB para consultar "
    "terabytes de Parquet sin servidor, cron/Airflow para programar. El SQL que "
    "escribieron se transfiere casi sin cambios.",
    "Convierte un notebook puntual en infraestructura — y tranquiliza a los "
    "alumnos de que las habilidades escalan a un proyecto real. Cambia el AOI y "
    "la base de datos se construye sola.",
    None)

# ── GLOSARIO ────────────────────────────────────────────────────────────
h1("Apéndice — Glosario en una línea (respuestas rápidas bajo presión)")
body("Definiciones concisas de los términos que más preguntan los alumnos. Cada "
     "una está formulada para que puedas decirla en voz alta tal cual.")

glossary = [
    ("Agua Case-2", "Agua ópticamente compleja donde el sedimento, el CDOM y/o la "
     "vegetación de fondo — no solo el fitoplancton — controlan el color, "
     "rompiendo los algoritmos oceánicos estándar."),
    ("C2RCC", "Corrección atmosférica con red neuronal entrenada para agua Case-2 "
     "que estima la reflectancia de teledetección robusta a esos confusores."),
    ("L1C vs L2A", "L1C es reflectancia tope de atmósfera; L2A es reflectancia de "
     "superficie tras la corrección atmosférica. Usa L2A para agua — pero recuerda "
     "que es una corrección de tierra, no C2RCC."),
    ("NDCI", "Índice de clorofila por diferencia normalizada = (B05 − B04)/(B05 + "
     "B04); usa el red-edge, así que sigue la clorofila incluso en agua turbia."),
    ("STAC", "SpatioTemporal Asset Catalog — el estándar moderno y programable "
     "para descubrir imágenes. La raíz actual de CDSE es "
     "stac.dataspace.copernicus.eu/v1."),
    ("COG", "Cloud-Optimised GeoTIFF — en mosaicos internos para leer solo la "
     "ventana del AOI por HTTP en vez de la escena entera."),
    ("CDSE", "Copernicus Data Space Ecosystem — el punto de acceso gratuito actual "
     "a los datos Sentinel, que reemplaza al retirado SciHub."),
    ("Process API", "Servicio de Sentinel Hub que ejecuta tu evalscript en "
     "servidor y devuelve el producto terminado para cualquier fecha."),
    ("Statistical API", "Servicio de Sentinel Hub que devuelve estadísticas sobre "
     "un AOI para todo un rango temporal en una petición — sin descargar imágenes."),
    ("TimeSeriesSplit", "Validación cruzada que respeta el orden cronológico; "
     "responde «¿puedo predecir fechas futuras?»."),
    ("GroupKFold (por estación)", "Validación cruzada que deja fuera estaciones "
     "enteras; responde la más difícil «¿puedo predecir un sitio no visto?»."),
    ("Isolation Forest", "Detector de anomalías no supervisado que aísla puntos "
     "multivariantes inusuales; lo entrenamos solo con años tranquilos pre-crisis."),
    ("Lead time vs retraso de detección", "Lead time = días que una marca salta "
     "ANTES de una crisis (acierto); retraso = días DESPUÉS (fallo). Nunca los "
     "mezcles."),
    ("Theil-Sen", "Estimador de tendencia robusto a valores atípicos — la opción "
     "correcta cuando unos picos de bloom distorsionarían una regresión ordinaria."),
    ("Escritura idempotente", "Una escritura que, repetida, deja la base de datos "
     "igual — aquí con UNIQUE + INSERT OR REPLACE — para que un trabajo "
     "programado nunca duplique."),
    ("SQLite + Parquet", "El almacenamiento híbrido: SQLite (un archivo, SQL "
     "completo) para tablas consultables; Parquet (columnar, comprimido) para los "
     "arrays masivos por píxel."),
    ("DANA", "Depresión Aislada en Niveles Altos — una gota fría que produce "
     "lluvia convectiva extrema sobre el SE de España; el disparador de la "
     "inundación de 2019."),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Light List Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Término"; hdr[1].text = "Di esto"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)
widths = (Inches(1.9), Inches(4.6))
for term, definition in glossary:
    cells = table.add_row().cells
    cells[0].text = term; cells[1].text = definition
    cells[0].paragraphs[0].runs[0].font.bold = True
    cells[0].paragraphs[0].runs[0].font.color.rgb = ACCENT
    for ci, w in enumerate(widths): cells[ci].width = w
for ci, w in enumerate(widths): table.rows[0].cells[ci].width = w

doc.add_paragraph()
foot = doc.add_paragraph()
rf = foot.add_run("Taller EO Mar Menor — guía del instructor. Combínala con las "
                  "presentaciones (slides/) y los tres notebooks (notebooks/).")
rf.italic = True; rf.font.size = Pt(9.5); rf.font.color.rgb = GREY

doc.save(OUT)
print("Guía de estudio (ES) escrita en", OUT)
