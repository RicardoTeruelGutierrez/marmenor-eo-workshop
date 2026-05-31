"""Build the instructor study guide / teaching script (.docx).

A combined, section-by-section guide across all three notebooks: for each part,
"in one line", "what to say", "why it matters", and "likely questions".
Run:  python scripts/build_study_guide.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "Instructor_Study_Guide.docx"

INK   = RGBColor(0x0E, 0x1A, 0x26)
ACCENT= RGBColor(0xC2, 0x41, 0x0C)
TEAL  = RGBColor(0x11, 0x5E, 0x59)
GREY  = RGBColor(0x6B, 0x61, 0x57)

doc = Document()

# ── Base styles ──────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for sec in doc.sections:
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)


def _style_heading(level, size, color, before, after):
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Georgia"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

_style_heading(1, 18, ACCENT, 18, 8)
_style_heading(2, 13.5, TEAL, 14, 4)
_style_heading(3, 11.5, INK, 8, 2)


def h1(text, page_break=True):
    p = doc.add_heading(text, level=1)
    if page_break:
        p.paragraph_format.page_break_before = True
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def body(text, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color is not None: r.font.color.rgb = color
    if size is not None: r.font.size = Pt(size)
    return p

def labeled(label, text, label_color=ACCENT):
    p = doc.add_paragraph()
    lr = p.add_run(label + " ")
    lr.bold = True; lr.font.color.rgb = label_color
    p.add_run(text)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead + " "); r.bold = True
    p.add_run(text)
    return p

def qa(pairs):
    p = doc.add_paragraph(); r = p.add_run("Likely questions"); r.bold = True; r.font.color.rgb = TEAL
    for q, a in pairs:
        pq = doc.add_paragraph()
        pq.paragraph_format.space_after = Pt(2)
        rq = pq.add_run("Q.  "); rq.bold = True; rq.font.color.rgb = ACCENT
        rq2 = pq.add_run(q); rq2.bold = True
        pa = doc.add_paragraph()
        pa.paragraph_format.left_indent = Inches(0.3)
        pa.paragraph_format.space_after = Pt(8)
        ra = pa.add_run("A.  "); ra.bold = True
        pa.add_run(a)


def section(title, oneline, say, why, questions=None):
    h2(title)
    labeled("In one line:", oneline)
    labeled("What to say:", say, label_color=TEAL)
    labeled("Why it matters:", why, label_color=TEAL)
    if questions:
        qa(questions)


# ===========================================================================
# TITLE
# ===========================================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("Mar Menor — Earth Observation Workshop")
r.font.name = "Georgia"; r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = INK
sub = doc.add_paragraph()
rs = sub.add_run("Instructor Study Guide & Teaching Script")
rs.font.name = "Georgia"; rs.font.size = Pt(16); rs.italic = True; rs.font.color.rgb = ACCENT
meta = doc.add_paragraph()
rm = meta.add_run("What to say, why it matters, and the questions you'll get — "
                  "section by section across all three notebooks.")
rm.font.size = Pt(12); rm.font.color.rgb = GREY

doc.add_paragraph()
labeled("How to use this guide:",
        "Each notebook section below has four parts — a one-line summary, a "
        "ready-to-speak explanation, the reason it matters, and the questions "
        "students tend to ask with model answers. Read it once end-to-end to own "
        "the narrative; skim the bold labels in the room. Pair it with the slide "
        "decks (slides/module1–3) and the live notebooks.")

labeled("The through-line:", "Three notebooks, one lagoon. Module 1 reads the "
        "lagoon from space; Module 2 validates that against the buoys and adds "
        "machine learning; Module 3 teaches students to build the whole data "
        "layer themselves. Everything centres on the Mar Menor's documented "
        "crises (2016, 2019 DANA, 2021 fish-kill, 2025 resurgence).")

# ===========================================================================
# PART 0
# ===========================================================================
h1("Part 0 — The story you are telling")
body("The Mar Menor is Europe's largest hypersaline coastal lagoon: ~135 km², "
     "max depth 7 m, separated from the Mediterranean by La Manga, a 22 km sand "
     "bar. Decades of nitrate runoff from the Campo de Cartagena agriculture, "
     "plus episodic DANA floods and rising Mediterranean temperatures, have "
     "tipped it into repeated eutrophication crises. Every collapse is visible "
     "from orbit at 10-metre resolution — which is exactly why it is a perfect "
     "teaching case.")
labeled("The arc to keep in your head:",
        "pressure (nutrients) → trigger (heat / flood) → bloom (chlorophyll) → "
        "oxygen collapse (anoxia) → fish-kill. Each module illuminates part of "
        "this chain with a different data source.")
labeled("One honest framing to repeat:",
        "Earth observation can shorten the loop from signal to decision, but the "
        "slow variable is the politics. Every algorithm here already runs in "
        "someone's pipeline in Murcia; the lagoon still degrades. Hold that "
        "tension — it keeps the technical work grounded.")

# ===========================================================================
# MODULE 1
# ===========================================================================
h1("Module 1 — Satellite Data Extraction")
body("Goal: from raw Sentinel reflectance to the visible fingerprints of the "
     "collapse. The notebook is split into three sections — A (synthetic, runs "
     "offline), B (the real CDSE/STAC workflow, shown not run), C (figures from "
     "real downloaded Sentinel-2 scenes).")

section("Setup, dependency check and the A/B/C split",
    "The notebook fails fast if a package or data file is missing, then declares "
    "its three-section structure.",
    "Start by pointing at the header table: Section A runs with no internet, "
    "Section B needs a free account, Section C needs the real GeoTIFFs (and "
    "self-skips if they're absent). Run the setup cell; it checks numpy/pandas/"
    "rasterio etc. and loads the lagoon outline.",
    "It directly answers a reviewer concern that synthetic and real data were "
    "mixed. The fail-fast checks mean a student never debugs a cryptic error "
    "halfway through.",
    [("Why not just one big notebook?",
      "Because 'what runs anywhere' and 'what needs the cloud' have different "
      "failure modes. Separating them keeps the class moving even with no Wi-Fi "
      "or no account.")])

section("A.1 — Case-2 waters: why standard algorithms fail",
    "The Mar Menor's optics are dominated by sediment, CDOM and macrophytes, not "
    "just phytoplankton — so open-ocean chlorophyll algorithms break.",
    "Explain Case-1 vs Case-2. Case-1 (open ocean) assumes phytoplankton alone "
    "drives the colour, so a simple blue/green ratio works. The Mar Menor is "
    "Case-2: suspended sediment raises blue-green reflectance, CDOM (coloured "
    "dissolved organics from runoff) absorbs blue, and bottom macrophytes add a "
    "near-infrared signal. Each moves the spectrum independently of chlorophyll. "
    "The fix is the C2RCC neural-network atmospheric correction, then a "
    "locally-calibrated index. Our proxy is the red-edge ratio Rrs(B05)/Rrs(B04).",
    "It is the conceptual backbone of the whole satellite story. If students "
    "miss this, they'll over-trust any single band ratio.",
    [("Is Sentinel-2 L2A the same as C2RCC?",
      "No — and say this clearly. L2A (Sen2Cor) is a LAND atmospheric correction. "
      "For rigorous water work you run C2RCC (e.g. in SNAP or via a Sentinel Hub "
      "evalscript). Our cached values only imitate a C2RCC-style product."),
     ("Why the red-edge (B05) and not classic blue/green?",
      "Chlorophyll absorbs strongly in the red (B04, 665 nm) and reflects in the "
      "red-edge (B05, 705 nm). The ratio rises with concentration and is far more "
      "robust to sediment than blue/green in turbid water.")])

section("A.3 — The lagoon-wide chlorophyll time series",
    "Averaging the sampling points to a weekly mean reproduces the four "
    "documented crises in the right months.",
    "Show the curve with the 2016, 2019, 2021 and 2025 windows shaded. Make the "
    "point that the synthetic data is calibrated against published literature, so "
    "the crises appear at the correct dates and magnitudes — it is a teaching "
    "stand-in for what a multi-year real pipeline would produce.",
    "It anchors the rest of the module in real, dated events and sets up 2021 as "
    "the focus crisis.",
    [("These are synthetic — does that undermine the lesson?",
      "No, as long as we're explicit. The statistics are calibrated to the "
      "literature; the methods (resampling, anomaly logic) are exactly what you'd "
      "run on real data.")])

section("A.6 — The 2021 marine heatwave and the oxygen link",
    "A +2 °C heatwave lowers oxygen solubility by ~0.3 mg/L; with the bloom "
    "respiring at night and stratification sealing the bottom, that drives anoxia.",
    "Walk the causal chain and keep it quantitative. August 2021 ran ~30.5 °C vs "
    "~27.5 °C in a normal year — a +1.9 °C anomaly. Using the Benson–Krause "
    "solubility curve, that warming alone removes ~0.3 mg/L of the water's oxygen "
    "capacity. At night the bloom consumes oxygen; stratification stops the "
    "bottom from re-oxygenating. Add it up and you get the hypoxia behind the "
    "fish-kill. We show the SST field clipped to the real lagoon outline, its "
    "anomaly, and the solubility curve side by side.",
    "It turns SST from a 'nice extra layer' into a computable driver of the "
    "ecological collapse — the emotional and scientific peak of Module 1.",
    [("Why clip the SST to the lagoon outline?",
      "Earlier versions showed pixels spilling onto land, which looked wrong and "
      "diluted the signal. Clipping to the NIR-derived outline keeps the stat "
      "honest and the figure clean."),
     ("Is 0.3 mg/L really enough to matter?",
      "On its own, no — but it stacks with night-time biological oxygen demand "
      "and stratification. The lesson is that several modest stressors combine "
      "non-linearly into a kill event.")])

section("Section B — The real CDSE / STAC / COG workflow",
    "How to discover and download real Sentinel data today, shown as patterns "
    "(not run in class).",
    "Emphasise three corrections to what students may have read in older "
    "tutorials. (1) SciHub is retired; the entry point is the Copernicus Data "
    "Space Ecosystem (CDSE). (2) The current STAC endpoint is "
    "stac.dataspace.copernicus.eu/v1 and the collection id is lowercase "
    "sentinel-2-l2a (the old catalogue endpoint was deprecated in Nov 2025). "
    "(3) Use L2A (surface reflectance), not L1C (top-of-atmosphere). Then the "
    "scalability trick: a Cloud-Optimised GeoTIFF lets you read only the AOI "
    "window — a few MB instead of an ~800 MB scene.",
    "These are the exact details that make student code work or fail in 2026. "
    "The COG window-read is the single most useful habit for scaling.",
    [("Why don't we run the authenticated calls in class?",
      "Credentials are personal and quota is limited. We show the pattern and run "
      "the no-login AWS mirror instead, so everyone succeeds.")])

section("Section C — Real Sentinel-2 figures (True Colour, NDCI, zones)",
    "Four real cloud-free scenes give pixel-perfect True-Colour, a per-pixel NDCI "
    "chlorophyll map, and a North/Centre/South zone analysis.",
    "Stress two methodological wins. (1) Scene selection: we measured valid "
    "coverage AND cloud over the lagoon itself, not the catalogue's whole-tile "
    "cloud %. That's why the 14-Jul-2021 scene was rejected — it only clipped a "
    "swath edge (5.5 % coverage). (2) The lagoon outline is extracted from the "
    "near-infrared band (water absorbs NIR), giving 191 exact vertices instead of "
    "a hand-drawn polygon. NDCI runs from ≈ −0.28 in clear winter water to ≈ "
    "−0.05 in the summer bloom; the zone bars show the north-west (agricultural) "
    "shore highest.",
    "This is where the abstract index becomes a believable map of a real place, "
    "and where students learn that data cleaning decisions change the science.",
    [("Why does the NDCI look mostly negative?",
      "Absolute NDCI depends on the atmospheric correction and the exact bands; "
      "what matters is the relative jump (~+0.25 from winter to bloom) and the "
      "spatial pattern, both of which are robust."),
     ("Could we compute real chlorophyll in mg/m³?",
      "Yes, with a locally-calibrated C2RCC + regression. We keep the real scenes "
      "as NDCI and use the synthetic series for the mg/m³ trend, to stay honest "
      "about what each data source supports.")])

# ===========================================================================
# MODULE 2
# ===========================================================================
h1("Module 2 — In-situ Integration & Machine Learning")
body("Goal: connect the satellite record to ground truth, validate it, and learn "
     "to predict and flag crises — with the statistical honesty that separates a "
     "teaching demo from an over-sold result.")

section("Section 1 — The in-situ monitoring networks",
    "Daily buoy data (CARM/IMIDA) plus watershed hydrology (CHS-SAIH) are the "
    "ground truth and the upstream pressure.",
    "Introduce the four sources: CARM real-time lagoon sondes (T, S, chl-a, DO, "
    "turbidity), IMIDA agro-meteorology (the nutrient pressure), CHS-SAIH "
    "watershed hydrology (the flood signal), and Copernicus Marine INS-TAC. Our "
    "cached file mimics CARM exports: five stations, daily, 2016–2025.",
    "Satellites need ground truth to be trusted; and the upstream agricultural "
    "data is what ultimately drives the lagoon.",
    None)

section("Section 1.1–1.2 — The crises in the raw signal",
    "Before any modelling, plot the 2021 hypoxia and the 2019 DANA directly from "
    "the buoys.",
    "For 2021, plot dissolved oxygen against the 2 mg/L hypoxia line with SST "
    "overlaid — this is the chart that accompanied the fish-kill headlines. For "
    "2019, show the DANA as a cascade: 300+ mm rain in 48 h → salinity drops "
    "~4 PSU → nitrate spikes ten-fold. ",
    "It builds the habit of eyeballing events before trusting any model on them, "
    "and it makes the later anomaly detection concrete.",
    None)

section("Section 2 — Match-ups: validating the satellite",
    "Co-locate each Sentinel-2 point with its nearest buoy (4 km, same day) and "
    "score the retrieval with proper metrics.",
    "Explain the match-up convention (ocean-colour uses 3 km / 3 h; we use 4 km / "
    "same-day because in-situ is daily). Report R², RMSE and RMSLE (log form, "
    "because chlorophyll is roughly log-normal) and bias, and always draw the 1:1 "
    "line. The scatter against 1:1 is the honest visual test.",
    "No skill claim is credible without a match-up. This is the discipline that "
    "separates a pretty map from a validated product.",
    [("Why log-transform the error (RMSLE)?",
      "Chlorophyll spans orders of magnitude and is approximately log-normal; an "
      "absolute RMSE is dominated by a few high-bloom points, hiding performance "
      "at typical concentrations.")])

section("Section 3 — ML retrieval and HONEST cross-validation",
    "A gradient-boosting regressor predicts chl-a from reflectances; we validate "
    "it two ways that answer different questions.",
    "First, TimeSeriesSplit honours chronological order and answers 'can we "
    "predict future dates?' — it looks strong. Then GroupKFold with the station "
    "as the group holds out ENTIRE stations and answers 'can we predict an "
    "unseen, unmonitored location?' — and the R² drops to about 0.17. The gap is "
    "the lesson: a model can look great in time yet be weak in space. Report the "
    "harder, spatial number when you talk about deploying somewhere new.",
    "This is the module's headline on rigor — it directly fixes the 'ML looks "
    "too optimistic' critique and teaches a transferable habit.",
    [("Why is the spatial R² so much lower — is the model bad?",
      "Not bad — honestly evaluated. Predicting a station the model never saw is "
      "genuinely harder than predicting a new date at a known station. The low "
      "number is the truthful deployment estimate."),
     ("Which number do I put in a paper?",
      "Both, clearly labelled. But the one that reflects real-world use at a new "
      "site is the GroupKFold (spatial) score.")])

section("Section 4 — Multivariate anomaly detection (early warning)",
    "An Isolation Forest trained ONLY on the quiet 2017–2018 years, then applied "
    "out-of-sample, flags joint anomalies across six variables.",
    "Make the training-window point central. If you fit the detector on the whole "
    "series, it learns the crises as 'normal' and the early-warning claim becomes "
    "circular. We fit on the calm 2017–2018 baseline and apply forward to 2019+, "
    "so every later flag is genuinely out-of-sample. It watches chl-a, turbidity, "
    "DO, SST, salinity and nitrate jointly — joint anomalies catch precursors a "
    "single chl-a threshold misses.",
    "It is the difference between a detector that 'predicts' the past and one that "
    "could actually warn an agency.",
    [("Why Isolation Forest and not a simple threshold?",
      "A single-variable threshold misses combinations — e.g. moderately high "
      "chl-a plus low DO plus high SST together. Isolation Forest scores the "
      "joint multivariate state.")])

section("Section 4.1 — Lead time vs detection delay",
    "Separate a flag BEFORE a crisis (lead time, a win) from a flag AFTER it "
    "(detection delay, a miss) — never conflate them.",
    "Show the table: 2019 was flagged 42 days early (a genuine warning); 2021 and "
    "2025 were only caught 1–2 days after onset (detection delay, i.e. misses). "
    "The naive code scanned a window that started before and ended after the "
    "event, took the first flag and called it 'lead time' even when it came late. "
    "We split the two metrics explicitly.",
    "It's a subtle bookkeeping fix that prevents monitoring papers from quietly "
    "overselling themselves — a great research-integrity teaching moment.",
    [("So our detector mostly fails?",
      "It succeeds spectacularly once (2019) and is late twice. That honest "
      "scorecard is the point — and it motivates adding watershed-forcing features "
      "to push the lead time earlier, which is the take-home exercise.")])

# ===========================================================================
# MODULE 3
# ===========================================================================
h1("Module 3 — Build Your Own EO Database")
body("Goal: students build the entire data layer themselves and run real "
     "downloads live. Frame the whole module with the kitchen metaphor: grocery "
     "run → wash → chop → fridge → cook → run the restaurant.")

section("Part 1 overview — the pipeline & the data landscape",
    "Six steps (discover, clean, extract, store, analyse, productionise) and a "
    "map of eight free data sources, four of them no-login.",
    "Give the whole map before the details. Then the rule of thumb: prototype "
    "today with the four no-login sources (NASA GIBS, AWS Sentinel-2 COGs, "
    "Open-Meteo Air-Quality and Marine); graduate to account-based ones (CDSE, "
    "Copernicus Marine, NASA Earthdata) when you need a specific product or the "
    "full archive.",
    "Students leave knowing not just one download trick but where to get almost "
    "any free environmental data — the real goal of the module.",
    None)

section("Section 1.1–1.4 — Live downloads (run these in the room)",
    "Four real, no-credentials downloads the students watch arrive: a GIBS image, "
    "a Sentinel-2 COG window, and two Open-Meteo time series.",
    "Run each cell live. GIBS: one URL returns a true-colour PNG for any date "
    "since 2000 — change the date and re-run. Sentinel-2 COG: search the "
    "catalogue, open the COG header WITHOUT downloading, then read only the AOI "
    "window — about 8 MB, ~3 % of the scene. Open-Meteo: a GET returns JSON that "
    "becomes a DataFrame of hourly PM2.5/PM10/ozone (pollution!) and "
    "sea-surface temperature. Each cell degrades gracefully if offline.",
    "This is the 'EO stops feeling abstract' moment — students see real bytes "
    "arrive and realise they can do it for their own coast today.",
    [("Is the COG window really downloading just part of the file?",
      "Yes. A Cloud-Optimised GeoTIFF is internally tiled and indexed; rasterio "
      "issues HTTP range requests for only the tiles overlapping your window. "
      "That's why it's ~3 % of the bytes."),
     ("Why use AWS instead of Copernicus for the live demo?",
      "The AWS Open-Data mirror needs no login, so every student succeeds "
      "instantly. Copernicus comes next, for what the account unlocks.")])

section("Section 1.5 — A free Copernicus account and what it unlocks",
    "Registration takes ~2 minutes; an OAuth client then enables server-side "
    "processing the AWS mirror can't do.",
    "Walk the three registration steps (register at dataspace.copernicus.eu → "
    "create an OAuth client in the Sentinel Hub dashboard → store id/secret as "
    "environment variables). Then the two unlocks, both verified live: the "
    "Process API sends a tiny evalscript and returns a finished product (NDCI "
    "map, true-colour image, or raw band stack) for any date; the Statistical API "
    "returns a full AOI time series in one request while downloading NO imagery. "
    "Show the spectral signature from the raw bands — reflectance rising into the "
    "red-edge is the chlorophyll fingerprint.",
    "It shows the leap from 'download pixels and compute yourself' to 'ask the "
    "server for exactly the product you want' — the professional workflow.",
    [("What does the account cost?",
      "Nothing — free, no credit card, ~30,000 processing units/month, plenty for "
      "a PhD project."),
     ("What is an evalscript?",
      "A small JavaScript function that tells Sentinel Hub what each output pixel "
      "should be. Change the function and you change the product — NDCI, true "
      "colour, or raw bands — without touching the download code.")])

section("Sections 2–4 — Clean, extract, and store (SQLite + Parquet)",
    "Turn pixels into trustworthy tidy rows, then store them in a typed, indexed, "
    "idempotent database plus Parquet for bulk arrays.",
    "Clean: mask non-water and clouds, score coverage. Extract: compute NDCI per "
    "scene and per zone into long-format rows (date, source, variable, zone, "
    "value). Store: a SQLite schema with primary keys, indices and — crucially — "
    "a UNIQUE constraint so INSERT OR REPLACE makes re-runs idempotent (no "
    "duplicates). Keep the millions of per-pixel values in Parquet, referenced "
    "from the database, not stuffed into SQL rows.",
    "This is the engineering heart: why a database beats a pile of CSVs, and why "
    "idempotency is what lets a pipeline run on a schedule forever.",
    [("Why SQLite and not PostgreSQL?",
      "SQLite is one file, zero install, full SQL — perfect for a PhD-scale "
      "project. We show the PostGIS/DuckDB upgrade path for when you outgrow it."),
     ("What does 'idempotent' buy me?",
      "You can re-process any date and the database updates that row instead of "
      "appending a duplicate — essential for a nightly job that might overlap "
      "previous runs.")])

section("Section 5 — The pollution study, straight from SQL",
    "Short SQL queries produce a monthly trend, a 'poor status' exceedance count, "
    "zone comparisons and a satellite↔in-situ join.",
    "Demonstrate that the questions a monitoring agency asks are now one query "
    "each: monthly-mean chlorophyll, a robust Theil-Sen trend (resistant to bloom "
    "spikes), the share of observations above a 'poor status' threshold per year, "
    "and a JOIN of satellite and in-situ. The exceedance counts spike in exactly "
    "2019 and 2021 — the database detects the documented crises by itself.",
    "It closes the loop: the data layer they built answers real pollution "
    "questions, proving the whole pipeline works end to end.",
    [("Why Theil-Sen instead of ordinary least squares?",
      "Theil-Sen is robust to outliers — a few extreme bloom months won't drag "
      "the trend line the way they would in OLS.")])

section("Section 6 — Productionise & scale",
    "Wrap steps 1–4 in one idempotent function, schedule it, and know the "
    "PostGIS / DuckDB upgrade path.",
    "Show the update_database() skeleton: discover → download window → clean → "
    "extract → upsert, all idempotent. Then the scaling table: PostGIS for "
    "spatial queries and many users, DuckDB for querying terabytes of Parquet "
    "with no server, cron/Airflow for scheduling. The SQL they wrote transfers "
    "almost unchanged.",
    "It turns a one-off notebook into infrastructure — and reassures students the "
    "skills scale to a real research project. Swap the AOI and the database "
    "builds itself.",
    None)

# ===========================================================================
# GLOSSARY
# ===========================================================================
h1("Appendix — One-line glossary (quick answers under pressure)")
body("Crisp definitions for the terms students ask about most. Each is phrased "
     "so you can say it out loud verbatim.")

glossary = [
    ("Case-2 water", "Optically complex water where sediment, CDOM and/or bottom "
     "vegetation — not just phytoplankton — drive the colour, breaking standard "
     "ocean-colour algorithms."),
    ("C2RCC", "A neural-network atmospheric correction trained for Case-2 water "
     "that estimates remote-sensing reflectance robust to those confounders."),
    ("L1C vs L2A", "L1C is top-of-atmosphere reflectance; L2A is surface "
     "reflectance after atmospheric correction. Use L2A for water — but remember "
     "it's a land correction, not C2RCC."),
    ("NDCI", "Normalized Difference Chlorophyll Index = (B05 − B04)/(B05 + B04); "
     "uses the red-edge, so it tracks chlorophyll even in turbid water."),
    ("STAC", "SpatioTemporal Asset Catalog — the modern, scriptable standard for "
     "discovering satellite imagery. CDSE's current root is "
     "stac.dataspace.copernicus.eu/v1."),
    ("COG", "Cloud-Optimised GeoTIFF — internally tiled so you can read just the "
     "AOI window over HTTP instead of the whole scene."),
    ("CDSE", "Copernicus Data Space Ecosystem — the current free access point for "
     "Sentinel data, replacing the retired SciHub."),
    ("Process API", "Sentinel Hub service that runs your evalscript server-side "
     "and returns the finished product for any date."),
    ("Statistical API", "Sentinel Hub service that returns statistics over an AOI "
     "for a whole time range in one request — no imagery downloaded."),
    ("TimeSeriesSplit", "Cross-validation that respects chronological order; "
     "answers 'can I predict future dates?'."),
    ("GroupKFold (by station)", "Cross-validation that holds out entire stations; "
     "answers the harder 'can I predict an unseen location?'."),
    ("Isolation Forest", "An unsupervised anomaly detector that isolates unusual "
     "multivariate points; we train it only on quiet pre-crisis years."),
    ("Lead time vs detection delay", "Lead time = days a flag fires BEFORE a "
     "crisis (a win); detection delay = days AFTER (a miss). Never merge them."),
    ("Theil-Sen", "A trend estimator robust to outliers — the right choice when a "
     "few bloom spikes would distort an ordinary regression."),
    ("Idempotent write", "A write that, repeated, leaves the database unchanged — "
     "here via UNIQUE + INSERT OR REPLACE — so a scheduled job never duplicates."),
    ("SQLite + Parquet", "The hybrid store: SQLite (one file, full SQL) for "
     "queryable tables; Parquet (columnar, compressed) for bulk per-pixel arrays."),
    ("DANA", "Depresión Aislada en Niveles Altos — a cut-off cold drop producing "
     "extreme convective rain over SE Spain; the 2019 flood trigger."),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Light List Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Term"; hdr[1].text = "Say this"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)
# set column widths
widths = (Inches(1.9), Inches(4.6))
for term, definition in glossary:
    cells = table.add_row().cells
    cells[0].text = term
    cells[1].text = definition
    cells[0].paragraphs[0].runs[0].font.bold = True
    cells[0].paragraphs[0].runs[0].font.color.rgb = ACCENT
    for ci, w in enumerate(widths):
        cells[ci].width = w
for ci, w in enumerate(widths):
    table.rows[0].cells[ci].width = w

doc.add_paragraph()
foot = doc.add_paragraph()
rf = foot.add_run("Mar Menor EO Workshop — instructor study guide. Pair with the "
                  "slide decks (slides/) and the three notebooks (notebooks/).")
rf.italic = True; rf.font.size = Pt(9.5); rf.font.color.rgb = GREY

doc.save(OUT)
print("Study guide written to", OUT)
print("Sections:", sum(1 for p in doc.paragraphs if p.style.name == 'Heading 2'))
